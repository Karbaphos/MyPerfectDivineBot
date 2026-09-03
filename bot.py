"""
Telegram-бот корпоративной базы знаний ИТ-службы.

Функциональность:
- при старте сканирует .md и .docx файлы в папке скрипта и разбивает их на пункты;
- поиск пункта по тексту запроса без БД и ML (стоп-слова + SequenceMatcher);
- ответы форматируются в HTML;
- отправка регламента (.md и .docx файлов) по команде.

Запуск: BOT_TOKEN=xxxx python bot.py
"""

import asyncio
import difflib
import html
import logging
import os
import re
from pathlib import Path
from typing import Optional

from aiogram import Bot, Dispatcher, F
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import Command, CommandStart
from aiogram.types import (
    FSInputFile,
    KeyboardButton,
    Message,
    ReplyKeyboardMarkup,
)
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

# ---------------------------------------------------------------------------
# Логирование
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger("kb_bot")

# ---------------------------------------------------------------------------
# Константы и настройки
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).resolve().parent

BTN_REGLAMENT = "📄 Прислать регламент"
BTN_HELP = "❓ Что умеет бот"
BTN_SEARCH = "🔍 Найти пункт"

MAIN_KEYBOARD = ReplyKeyboardMarkup(
    keyboard=[
        [KeyboardButton(text=BTN_REGLAMENT)],
        [KeyboardButton(text=BTN_HELP)],
        [KeyboardButton(text=BTN_SEARCH)],
    ],
    resize_keyboard=True,
)

MIN_COMMON_WORDS = 1
SCORE_THRESHOLD = 1.15

# Базовый набор русских стоп-слов
STOP_WORDS: set[str] = {
    "и", "в", "во", "не", "что", "он", "на", "я", "с", "со", "как", "а",
    "то", "все", "она", "так", "его", "но", "да", "ты", "к", "у", "же",
    "вы", "за", "бы", "по", "только", "ее", "мне", "было", "вот", "от",
    "меня", "еще", "нет", "о", "из", "ему", "теперь", "когда", "даже",
    "ну", "вдруг", "ли", "если", "уже", "или", "ни", "быть", "был",
    "него", "до", "вас", "нибудь", "опять", "уж", "вам", "ведь", "там",
    "потом", "себя", "ничего", "ей", "может", "они", "тут", "где",
    "есть", "надо", "ней", "для", "мы", "тебя", "их", "чем", "была",
    "сам", "чтоб", "без", "будто", "чего", "раз", "тоже", "себе",
    "под", "будет", "ж", "тогда", "кто", "этот", "того", "потому",
    "этого", "какой", "совсем", "ним", "здесь", "этом", "один",
    "почти", "мой", "тем", "чтобы", "нее", "сейчас", "были", "куда",
    "зачем", "всех", "никогда", "можно", "при", "наконец", "два",
    "об", "другой", "хоть", "после", "над", "больше", "тот", "через",
    "эти", "нас", "про", "всего", "них", "какая", "много", "разве",
    "три", "эту", "моя", "впрочем", "хорошо", "свою", "этой", "перед",
    "иногда", "лучше", "чуть", "том", "нельзя", "такой", "им", "более",
    "всегда", "конечно", "всю", "между", "это", "мои", "твой", "также",
    "которые", "который", "которая", "которое", "нужно", "можете",
    "пожалуйста", "здравствуйте", "привет", "подскажите", "скажите",
    "хочу", "хотел", "хотела", "нужен", "нужна", "нужны", "мне",
}

# ---------------------------------------------------------------------------
# Разбор .md файлов на пункты
# ---------------------------------------------------------------------------

# Ловим строки вида "1. Текст", "1.2 Текст", "# 1.2 Текст"
ITEM_PATTERN = re.compile(
    r"^(?:#{1,6}\s+)?(\d+(?:\.\d+)*)\.?\s+(\S.*)$",
    re.MULTILINE,
)


class KnowledgeItem:
    __slots__ = ("number", "title", "text", "source")

    def __init__(self, number: str, title: str, text: str, source: str):
        self.number = number
        self.title = title
        self.text = text
        self.source = source


def extract_text_from_md(path: Path) -> str:
    """Читает содержимое .md файла как есть."""
    return path.read_text(encoding="utf-8")


def extract_text_from_docx(path: Path) -> str:
    """Извлекает текст из .docx файла: параграфы (включая заголовки) и таблицы.

    ВАЖНО: номер пункта должен быть напечатан в самом тексте абзаца
    (например «1.2 Название»). Автоматическая нумерация списков Word
    (когда цифры генерируются самим Word и не хранятся как текст)
    не считывается — её нужно вписывать вручную.
    """
    document = DocxDocument(str(path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def split_into_items(content: str, source_name: str) -> list[KnowledgeItem]:
    """Разбивает произвольный текст на пункты по номерам ("1.", "1.2", "# 1.2")."""
    matches = list(ITEM_PATTERN.finditer(content))
    items: list[KnowledgeItem] = []

    for idx, match in enumerate(matches):
        number = match.group(1).strip()
        title = match.group(2).strip()
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(content)
        body = content[start:end].strip()
        full_text = f"{title}\n{body}".strip() if body else title

        items.append(
            KnowledgeItem(
                number=number,
                title=title,
                text=full_text,
                source=source_name,
            )
        )

    return items


def parse_source_file(path: Path) -> list[KnowledgeItem]:
    """Читает файл регламента (.md или .docx) и разбивает его на пункты."""
    suffix = path.suffix.lower()

    try:
        if suffix == ".md":
            content = extract_text_from_md(path)
        elif suffix == ".docx":
            content = extract_text_from_docx(path)
        else:
            logger.warning("Неподдерживаемый тип файла, пропускаю: %s", path)
            return []
    except PackageNotFoundError:
        logger.error(
            "Файл %s повреждён или не является настоящим .docx (возможно, это .doc)",
            path,
        )
        return []
    except Exception as exc:  # noqa: BLE001
        logger.error("Не удалось прочитать файл %s: %s", path, exc)
        return []

    items = split_into_items(content, path.name)
    logger.info("Файл %s: найдено пунктов — %d", path.name, len(items))
    return items


def load_knowledge_base(directory: Path) -> tuple[list[KnowledgeItem], list[Path]]:
    """Ищет все .md и .docx файлы в папке и разбирает их на пункты."""
    source_files = sorted(directory.glob("*.md")) + sorted(directory.glob("*.docx"))
    # Исключаем временные файлы Word (~$файл.docx), создаваемые при открытом документе
    source_files = [f for f in source_files if not f.name.startswith("~$")]

    if not source_files:
        logger.warning(
            "В папке %s не найдено ни одного .md или .docx файла", directory
        )

    all_items: list[KnowledgeItem] = []
    for source_file in source_files:
        all_items.extend(parse_source_file(source_file))

    logger.info(
        "База знаний загружена: файлов — %d, пунктов — %d",
        len(source_files), len(all_items),
    )
    return all_items, source_files


KNOWLEDGE_ITEMS: list[KnowledgeItem] = []
SOURCE_FILES: list[Path] = []

# ---------------------------------------------------------------------------
# Поиск пункта (без БД и ML)
# ---------------------------------------------------------------------------

_WORD_CLEAN_RE = re.compile(r"^[^\wа-яА-ЯёЁ]+|[^\wа-яА-ЯёЁ]+$")


def _normalize_word(word: str) -> str:
    return _WORD_CLEAN_RE.sub("", word).lower()


def _significant_words(text: str) -> set[str]:
    words = (_normalize_word(w) for w in text.split())
    return {w for w in words if w and w not in STOP_WORDS}


def find_item(query: str, items: list[KnowledgeItem]) -> Optional[KnowledgeItem]:
    """Ищет наиболее подходящий пункт по совпадению значимых слов и схожести текста."""
    query_words = _significant_words(query)
    if not query_words:
        return None

    query_lower = query.lower()
    best_item: Optional[KnowledgeItem] = None
    best_score = 0.0

    for item in items:
        item_words = _significant_words(f"{item.title} {item.text}")
        common_words = query_words & item_words

        if len(common_words) < MIN_COMMON_WORDS:
            continue

        similarity = difflib.SequenceMatcher(
            None, query_lower, item.text.lower()
        ).ratio()

        score = len(common_words) + similarity

        if score > SCORE_THRESHOLD and score > best_score:
            best_score = score
            best_item = item

    return best_item

# ---------------------------------------------------------------------------
# Форматирование ответа
# ---------------------------------------------------------------------------

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def markdown_to_html(text: str) -> str:
    """Безопасно экранирует текст и конвертирует базовую Markdown-разметку в HTML."""
    escaped = html.escape(text)
    escaped = _BOLD_RE.sub(r"<b>\1</b>", escaped)
    escaped = _ITALIC_RE.sub(r"<i>\1</i>", escaped)
    return escaped


def format_answer(item: KnowledgeItem) -> str:
    divider = "━" * 15
    body_html = markdown_to_html(item.text)
    return (
        f"📄 Документ: {html.escape(item.source)}\n"
        f"📎 Пункт {html.escape(item.number)}\n"
        f"{divider}\n"
        f"{body_html}"
    )

# ---------------------------------------------------------------------------
# Тексты сообщений
# ---------------------------------------------------------------------------

WELCOME_TEXT = (
    "👋 Здравствуйте!\n\n"
    "Я бот корпоративной базы знаний ИТ-службы.\n"
    "Я помогу быстро найти нужный пункт регламента — просто напишите "
    "вопрос своими словами.\n\n"
    "Выберите действие на клавиатуре ниже или напишите запрос."
)

HELP_TEXT = (
    "❓ <b>Что я умею</b>\n\n"
    "• /start — приветствие и главное меню\n"
    "• /help — эта справка\n"
    "• /reglament — пришлю файлы регламента целиком (.md и .docx)\n"
    "• Любой другой текст — найду подходящий пункт регламента "
    "по смыслу вашего вопроса\n\n"
    "Если пункт не найден, я так и сообщу и не буду ничего придумывать."
)

NOT_FOUND_TEXT = (
    "🤷 К сожалению, по вашему запросу подходящий пункт регламента не найден.\n"
    "Попробуйте переформулировать вопрос или используйте команду /reglament, "
    "чтобы посмотреть документы целиком."
)

NO_FILES_TEXT = (
    "⚠️ Файлы регламента (.md или .docx) пока не найдены рядом со скриптом бота.\n"
    "Обратитесь к администратору базы знаний."
)

ASK_QUERY_TEXT = "🔍 Напишите вопрос или ключевые слова — я найду подходящий пункт."

# ---------------------------------------------------------------------------
# Инициализация бота
# ---------------------------------------------------------------------------

BOT_TOKEN = os.environ.get("BOT_TOKEN")

dp = Dispatcher()


@dp.message(CommandStart())
async def handle_start(message: Message) -> None:
    logger.info("Пользователь %s вызвал /start", message.from_user.id if message.from_user else "?")
    await message.answer(WELCOME_TEXT, reply_markup=MAIN_KEYBOARD)


@dp.message(Command("help"))
@dp.message(F.text == BTN_HELP)
async def handle_help(message: Message) -> None:
    logger.info("Пользователь %s запросил справку", message.from_user.id if message.from_user else "?")
    await message.answer(HELP_TEXT, reply_markup=MAIN_KEYBOARD)


@dp.message(Command("reglament"))
@dp.message(F.text == BTN_REGLAMENT)
async def handle_reglament(message: Message) -> None:
    user_id = message.from_user.id if message.from_user else "?"
    logger.info("Пользователь %s запросил регламент", user_id)

    if not SOURCE_FILES:
        await message.answer(NO_FILES_TEXT, reply_markup=MAIN_KEYBOARD)
        return

    for source_file in SOURCE_FILES:
        try:
            document = FSInputFile(source_file)
            await message.answer_document(document)
        except Exception as exc:  # noqa: BLE001
            logger.error("Не удалось отправить файл %s пользователю %s: %s", source_file, user_id, exc)
            await message.answer(
                f"⚠️ Не удалось отправить файл {html.escape(source_file.name)}."
            )

    await message.answer("Это все файлы регламента.", reply_markup=MAIN_KEYBOARD)


@dp.message(F.text == BTN_SEARCH)
async def handle_search_button(message: Message) -> None:
    await message.answer(ASK_QUERY_TEXT, reply_markup=MAIN_KEYBOARD)


@dp.message(F.text)
async def handle_free_text(message: Message) -> None:
    query = (message.text or "").strip()
    user_id = message.from_user.id if message.from_user else "?"

    if not query:
        return

    logger.info("Пользователь %s ищет: %s", user_id, query)

    if not KNOWLEDGE_ITEMS:
        await message.answer(NO_FILES_TEXT, reply_markup=MAIN_KEYBOARD)
        return

    item = find_item(query, KNOWLEDGE_ITEMS)

    if item is None:
        logger.info("По запросу пользователя %s ничего не найдено", user_id)
        await message.answer(NOT_FOUND_TEXT, reply_markup=MAIN_KEYBOARD)
        return

    logger.info(
        "Пользователю %s найден пункт %s из %s", user_id, item.number, item.source
    )
    await message.answer(format_answer(item), reply_markup=MAIN_KEYBOARD)


async def main() -> None:
    if not BOT_TOKEN:
        logger.error(
            "Переменная окружения BOT_TOKEN не установлена. "
            "Установите её перед запуском: export BOT_TOKEN=ваш_токен"
        )
        raise SystemExit(1)

    global KNOWLEDGE_ITEMS, SOURCE_FILES
    KNOWLEDGE_ITEMS, SOURCE_FILES = load_knowledge_base(SCRIPT_DIR)

    bot = Bot(
        token=BOT_TOKEN,
        default=DefaultBotProperties(parse_mode=ParseMode.HTML),
    )

    logger.info("Бот запускается в режиме polling...")
    try:
        await dp.start_polling(bot)
    except Exception as exc:  # noqa: BLE001
        logger.exception("Критическая ошибка при работе бота: %s", exc)
        raise
    finally:
        await bot.session.close()


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except (KeyboardInterrupt, SystemExit):
        logger.info("Бот остановлен.")
